#coding:utf-8
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)  #插入标题

p = document.add_paragraph('A plain paragraph having some 1231312312312312312312312312312 3123123123123123123123123123123123123123123123')   #插入段落

table = document.add_table(rows=2, cols=4) #插入表格
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
hdr_cells[3].text = 'wwww'



document.add_page_break()

document.save('demo.docx')  #保存文档